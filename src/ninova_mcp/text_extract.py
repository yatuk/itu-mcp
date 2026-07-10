"""Extract plain text from downloaded or in-memory office/PDF files."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any
from urllib.parse import urlparse


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm"}
DEFAULT_MAX_CHARS = 50_000


def _extension_from_name(name: str | None) -> str:
    if not name:
        return ""
    return Path(name).suffix.lower()


def guess_extension(
    *,
    path: str | Path | None = None,
    url: str | None = None,
    content_type: str | None = None,
    filename: str | None = None,
) -> str:
    for candidate in (filename, str(path) if path else None):
        ext = _extension_from_name(candidate)
        if ext:
            return ext
    if url:
        ext = _extension_from_name(urlparse(url).path)
        if ext:
            return ext
    ctype = (content_type or "").split(";", 1)[0].strip().lower()
    mapping = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/plain": ".txt",
        "text/markdown": ".md",
        "text/csv": ".csv",
        "application/json": ".json",
        "text/html": ".html",
        "application/xml": ".xml",
        "text/xml": ".xml",
    }
    return mapping.get(ctype, "")


def _truncate(text: str, max_chars: int) -> tuple[str, bool]:
    if max_chars <= 0 or len(text) <= max_chars:
        return text, False
    return text[:max_chars], True


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise RuntimeError(
            "PDF extraction requires pypdf. Install with: pip install pypdf"
        ) from exc

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        page_text = page_text.strip()
        if page_text:
            parts.append(f"--- page {index} ---\n{page_text}")
    return "\n\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "DOCX extraction requires python-docx. Install with: pip install python-docx"
        ) from exc

    document = Document(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    # Tables often hold assignment rubrics / schedules.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs).strip()


def _extract_plain(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1254", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text_from_bytes(
    data: bytes,
    *,
    extension: str,
    max_chars: int = DEFAULT_MAX_CHARS,
) -> dict[str, Any]:
    ext = extension.lower() if extension.startswith(".") else f".{extension.lower()}" if extension else ""
    if ext not in SUPPORTED_EXTENSIONS and ext not in {".htm"}:
        return {
            "ok": False,
            "extension": ext or None,
            "error": (
                f"Unsupported file type {ext or '(unknown)'}. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            ),
            "text": None,
            "char_count": 0,
            "truncated": False,
        }

    try:
        if ext == ".pdf":
            text = _extract_pdf(data)
        elif ext == ".docx":
            text = _extract_docx(data)
        else:
            text = _extract_plain(data)
            if ext in {".html", ".htm"}:
                # Lightweight strip for HTML dumps without pulling BS4 into this path always.
                text = re.sub(r"(?is)<script[^>]*>.*?</script>", " ", text)
                text = re.sub(r"(?is)<style[^>]*>.*?</style>", " ", text)
                text = re.sub(r"(?s)<[^>]+>", " ", text)
                text = re.sub(r"\s+", " ", text).strip()
    except Exception as exc:
        return {
            "ok": False,
            "extension": ext,
            "error": str(exc),
            "text": None,
            "char_count": 0,
            "truncated": False,
        }

    text = (text or "").strip()
    truncated_text, truncated = _truncate(text, max_chars)
    return {
        "ok": True,
        "extension": ext,
        "error": None,
        "text": truncated_text,
        "char_count": len(text),
        "returned_chars": len(truncated_text),
        "truncated": truncated,
        "empty": not bool(text),
    }


def extract_text_from_path(
    path: str | Path,
    *,
    max_chars: int = DEFAULT_MAX_CHARS,
    extension: str | None = None,
) -> dict[str, Any]:
    file_path = Path(path).expanduser().resolve()
    if not file_path.is_file():
        return {
            "ok": False,
            "path": str(file_path),
            "error": f"File not found: {file_path}",
            "text": None,
            "char_count": 0,
            "truncated": False,
        }
    data = file_path.read_bytes()
    ext = extension or guess_extension(path=file_path)
    result = extract_text_from_bytes(data, extension=ext, max_chars=max_chars)
    result["path"] = str(file_path)
    result["size_bytes"] = len(data)
    return result
